import requests
from wordnik import *
import time
from urllib.error import HTTPError
from docx import Document

apiUrl = 'http://api.wordnik.com/v4'
apiKey = "h1d2viwpqzcb0ni0d1109aa673j5fgwgokvtq4gjuwykcl3l6"
client = swagger.ApiClient(apiKey, apiUrl)
wordApi = WordApi.WordApi(client)
word_doc = "C:/Users/prana/Downloads/Vocabulary_Assignment_TEMPLATE_Quick_Review_15_words(5).docx"
document = Document(word_doc)

userWord = ["acrimony", "balk", "cajole", "dour", "expound", "exult", "feasible", "fiasco", "fluctuate", "harry",
            "incognito", "inscrutable", "lethargy", "metier", "omniscient"]
x = 0
row = 3

while x in range(len(userWord)) and row < 36:
    nouns, adverbs, adj, verbs, syn, ant = {}, {}, {}, {}, [], []
    nouns, adverbs, adj, verbs = set(nouns), set(adverbs), set(adj), set(verbs)


    def categorize_words(word_list):
        for word in word_list:
            suffix = word[-2:len(word)]
            if suffix == "on" or word[-4:len(word)] == "ness" or word[-4:len(word)] == "ment" or word[-6:len(word)] == "nesses" or suffix == "rs" or suffix == "ry" or suffix == 'sm' or suffix == "ty" or suffix == "ge" or suffix == "ns" or word[3:len(word)] == "ies" or suffix == "ts" or suffix == "er":
                nouns.add(word)
            elif suffix == "ly":
                adverbs.add(word)
            elif suffix == "ic" or suffix == "nt" or suffix == "us" or suffix == "ve" or suffix == "al" or suffix == "le" or suffix == "st":
                adj.add(word)
            elif suffix == "ng" or suffix == "ed" or suffix == "te" or suffix == "ze" or suffix == "se" or suffix == "es":
                verbs.add(word)


    userWord[x] = userWord[x].lower()
    try:
        definition = wordApi.getDefinitions(userWord[x])
        wordForms = wordApi.getRelatedWords(userWord[x])
    except HTTPError:
        time.sleep(5)
        x = x
        continue

    for i in wordForms:
        if i.relationshipType == "form" or i.relationshipType == "stem" or i.relationshipType == "verb-form" or i.relationshipType == "verb-stem" or i.relationshipType == "inflected-form" or i.relationshipType == "etymologically-related-term":
            categorize_words(i.words)

        if i.relationshipType == "synonym":
            maxidx = 3
            if len(i.words) >= maxidx:
                for j in range(maxidx):
                    syn.append(i.words[j])
            else:
                maxidx = len(i.words)
                for j in range(maxidx):
                    syn.append(i.words[j])
    mer_url = 'https://dictionaryapi.com/api/v3/references/thesaurus/json/' + userWord[
        x] + '?key=022c0b10-4019-4ac1-b6ee-3f136afd1c67'
    mer_dic = 'https://dictionaryapi.com/api/v3/references/collegiate/json/' + userWord[
        x] + '?key=c2441927-e822-481e-9806-fee699bf54ee'
    r1 = requests.get(mer_url).json()
    r2 = requests.get(mer_dic)
    r2 = r2.json()
    categorize_words(r1[0]["meta"]["stems"])

    for i in r1:
        ants1 = i["meta"]["ants"]
        if not len(ants1) == 0:
            ant.append(i["meta"]['ants'][0][0])

    defintion_new = definition[0].text
    try:
        defintion_new = defintion_new.replace('<i>', "").replace('</i>', "").replace('<', "").replace('>', "").replace(
            '/', "")
    except AttributeError:
        defintion_new = r2[0]["shortdef"][0]

    wordCats, syn_ants = {}, {}
    wordCats["nouns"], wordCats["adverbs"], wordCats["adjectives"], wordCats["verbs"], syn_ants["synonyms"], syn_ants[
        "antonyms"] = list(nouns), list(adverbs), list(adj), list(verbs), syn, ant
    document.tables[0].cell(row, 0).add_paragraph(userWord[x] + ": " + defintion_new)
    for cat in wordCats:
        if len(wordCats[cat]) != 0:
            wordCats[cat] = str(wordCats[cat])
            document.tables[0].cell(row, 1).add_paragraph(
                cat + ": " + wordCats[cat].replace("[", "").replace("]", "").replace("'", ""))
        else:
            document.tables[0].cell(row, 1).add_paragraph(cat + ": " + "none")
    for nym in syn_ants:
        if len(syn_ants[nym]) != 0:
            syn_ants[nym] = str(syn_ants[nym])
            document.tables[0].cell(row, 2).add_paragraph(
                nym + ": " + syn_ants[nym].replace("[", "").replace("]", "").replace("'", ""))
        else:
            document.tables[0].cell(row, 2).add_paragraph(nym + ": " + "none")
    x += 1
    row += 2
document.save(word_doc)
